from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT_DIR / 'documentation' / 'word' / 'reference.docx'
EAST_ASIA_FONT = '游ゴシック'


def set_font(run, name='Yu Gothic', size=None, bold=None, color=None):
    if name:
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


doc = Document()

# Modify Normal style
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
style.font.size = Pt(10.5)

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Yu Gothic'
title_style._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
title_style.font.size = Pt(28)
title_style.font.bold = True

doc.styles['Subtitle'].font.name = 'Yu Gothic'
doc.styles['Subtitle']._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
doc.styles['Subtitle'].font.size = Pt(14)
doc.styles['Subtitle'].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for level in [1, 2, 3]:
    heading = doc.styles[f'Heading {level}']
    heading.font.name = 'Yu Gothic'
    heading._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_ASIA_FONT)
    heading.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    heading.font.bold = True
    heading.font.size = Pt(20 - (level - 1) * 2)

# Cover page
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('レポートタイトル (サンプル)')
set_font(run, size=Pt(32), bold=True)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('サブタイトルはサンプル情報としてここに記載します')
set_font(run, size=Pt(16), color=RGBColor(0x55, 0x55, 0x55))

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('作成日: 2025-09-29 / 作成者: Sample User')
set_font(run, size=Pt(13), color=RGBColor(0x55, 0x55, 0x55))

doc.add_page_break()

# Body sections
h1 = doc.add_heading('1. 概要', level=1)
para = doc.add_paragraph('本テンプレートは Pandoc の reference.docx として想定しています。')
para.add_run(' スタイルを調整することで、Markdown 変換後の Word 文書の整形を容易にします。')

h2 = doc.add_heading('1.1 使用するスタイル', level=2)
doc.add_paragraph('必要なスタイル一覧', style='List Bullet')
doc.add_paragraph('必要なスタイルは段階的にインデントされます。', style='List Bullet 2')

h2 = doc.add_heading('1.2 設定値', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
row = table.rows[0].cells
row[0].text = '項目'
row[1].text = '設定内容'
row[2].text = '備考'
row = table.add_row().cells
row[0].text = 'サンプル'
row[1].text = 'Markdown 変換後整形'
row[2].text = '手動微調整'
row = table.add_row().cells
row[0].text = 'フォント'
row[1].text = '游ゴシック / Yu Gothic'
row[2].text = '標準書体'

h2 = doc.add_heading('1.3 注意点', level=2)
para = doc.add_paragraph()
run = para.add_run('本文テキストはボールドやイタリックで強調可能です。')
run.bold = True
doc.add_paragraph('注意の例: この箇所はブロック引用を活用します。', style='Intense Quote')

h1 = doc.add_heading('2. まとめ', level=1)
doc.add_paragraph('本 reference.docx をプロジェクト毎にカスタマイズすることで、Markdown 変換後の書式差異を抑制できます。必要に応じて游ゴシック設定を微調整してください。')

for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = 'Confidential - Sample Reference Template'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f'reference.docx created at {OUTPUT_PATH}')
