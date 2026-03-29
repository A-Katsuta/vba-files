from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

OUTPUT_PATH = r"C:\github-vba-lab\github-vba-lab\reference.docx"


def set_font(r, name='Yu Gothic', size=None, bold=None, color=None):
    if name:
        r.font.name = name
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
    if size:
        r.font.size = size
    if bold is not None:
        r.font.bold = bold
    if color:
        r.font.color.rgb = color


doc = Document()

# Modify Normal style
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
style.font.size = Pt(10.5)

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Yu Gothic'
title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
title_style.font.size = Pt(28)
title_style.font.bold = True

doc.styles['Subtitle'].font.name = 'Yu Gothic'
doc.styles['Subtitle']._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
doc.styles['Subtitle'].font.size = Pt(14)
doc.styles['Subtitle'].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for level in [1, 2, 3]:
    heading = doc.styles[f'Heading {level}']
    heading.font.name = 'Yu Gothic'
    heading._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')
    heading.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    heading.font.bold = True
    heading.font.size = Pt(20 - (level-1)*2)

# Cover page
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('レポートタイトル (サンプル)')
set_font(run, size=Pt(32), bold=True)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('サブタイトルやバージョン情報をここに記載します')
set_font(run, size=Pt(16), color=RGBColor(0x55, 0x55, 0x55))

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('作成日: 2025-09-29 / 作成者: Sample User')
set_font(run, size=Pt(13), color=RGBColor(0x55, 0x55, 0x55))

doc.add_page_break()

# Body sections
h1 = doc.add_heading('1. 概要', level=1)
para = doc.add_paragraph('本テンプレートは Pandoc の reference.docx として利用することを想定しています。')
para.add_run(' スタイルを調整することで、Markdown から生成する Word 文書の統一感を高められます。')

h2 = doc.add_heading('1.1 箇条書きスタイル', level=2)
doc.add_paragraph('第一階層の箇条書き項目', style='List Bullet')
doc.add_paragraph('第二階層の箇条書きはインデントで表現します。', style='List Bullet 2')

h2 = doc.add_heading('1.2 表スタイル', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
row = table.rows[0].cells
row[0].text = '項目'
row[1].text = '内容'
row[2].text = '備考'
row = table.add_row().cells
row[0].text = 'サンプル'
row[1].text = 'Markdown から自動変換'
row[2].text = '参照ドキュメント'
row = table.add_row().cells
row[0].text = 'フォント'
row[1].text = '游ゴシック / Yu Gothic'
row[2].text = '日本語を想定'

h2 = doc.add_heading('1.3 引用・強調', level=2)
para = doc.add_paragraph()
run = para.add_run('強調テキストはボールドやイタリックで表現されます。')
run.bold = True
doc.add_paragraph('引用の例: 体験価値はブランド戦略に直結する。', style='Intense Quote')

h1 = doc.add_heading('2. まとめ', level=1)
doc.add_paragraph('この reference.docx をプロジェクト専用にカスタマイズすることで、Markdown から生成される Word 文書の品質と統一感を高められます。必要に応じてロゴの挿入や表紙デザインの調整を行ってください。')

for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = 'Confidential - Sample Reference Template'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"reference.docx created at {OUTPUT_PATH}")
